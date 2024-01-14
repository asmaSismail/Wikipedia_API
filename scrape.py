import requests 
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt ,RGBColor
import re 
#GET search title
url =  "https://en.wikipedia.org/wiki/"
print("Write Title : ")
search = input()
try :
    search= search.replace(' ','_')
    #fetch url
    result = requests.get( f'{url}{search}')
    #convert result to html parser
    soup = BeautifulSoup(result.text,'html.parser')
    #get only content section
    content = soup.find("div", {"id": "mw-content-text"})
    doc = Document()
    for element in content.find_all():
        #set text and color of word doc
        text_size = Pt(14)
        text_color = RGBColor(0, 0, 0)
        text = element.get_text().strip()
        #if title set new font
        if element.name not in ['h1','h2','h3','p']:
            pass
        else :
            if text in  ['See also','References','Notes']:
                break
            if element.name in ['h2','h1','h3']:
                text_size = Pt(18)  
                text_color= RGBColor(0, 0, 255)
            #delete references numbers
            text= re.sub(r'\[.*?\]', '',text)
            run = doc.add_paragraph().add_run(text)
            font = run.font
            font.size = text_size
            font.color.rgb= text_color
    #save resuts
    doc.save(f"{search}.docx")
except:
    print('PAGE NOT FOUND')